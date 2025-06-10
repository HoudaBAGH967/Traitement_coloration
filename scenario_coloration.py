#!/usr/bin/env python
# coding: utf-8

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
from docx.enum.text import WD_BREAK
import sys


fichier_data = sys.argv[1].strip()
fichier_output = sys.argv[2].strip()

# Couleurs associées aux feuilles Excel
feuilles_couleurs = {
    'CHAMPS_CRE': RGBColor(0x00, 0xB0, 0x50),        # Vert
    'CONNECTEUR_LOGIQUE': RGBColor(0xFF, 0x00, 0x00) # Rouge
}

# Préparer les mots-clés à colorer
mots_colores = {}
for feuille, couleur in feuilles_couleurs.items():
    try:
        df = pd.read_excel(fichier_data, sheet_name=feuille)
        if feuille == 'CHAMPS_CRE':
            for col in ['nom_champ', 'libelle_champ']:
                if col in df.columns:
                    mots = df[col].dropna().astype(str).str.strip().tolist()
                    for mot in mots:
                        mots_colores[mot.upper()] = couleur
        else:
            mots = df.iloc[:, 0].dropna().astype(str).str.strip().tolist()
            for mot in mots:
                mots_colores[mot.upper()] = couleur
    except Exception as e:
        print(f"Erreur lors du traitement de la feuille {feuille} : {e}")

# Regex sécurisée pour trouver les mots clés, insensible à la casse
escaped_keywords = sorted([re.escape(k) for k in mots_colores.keys()], key=len, reverse=True)
pattern = re.compile(r'(?<!\w)(' + '|'.join(escaped_keywords) + r')(?!\w)', re.IGNORECASE)

def colorer_texte(para):
    texte = para.text
    if not texte.strip():
        return
    para.clear()
    pos = 0
    for match in pattern.finditer(texte):
        if match.start() > pos:
            para.add_run(texte[pos:match.start()])
        run = para.add_run(match.group())
        run.font.color.rgb = mots_colores.get(match.group().upper(), RGBColor(0, 0, 0))
        run.bold = True
        pos = match.end()
    if pos < len(texte):
        para.add_run(texte[pos:])

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs[edge]
            tag = f'w:{edge}'
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))

def traiter_cellules(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                colorer_texte(para)
                for run in para.runs:
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                    run.font.size = Pt(8)
            for nested_table in cell.tables:
                traiter_cellules(nested_table)

# Création du document Word avec marges personnalisées
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)

# Chargement des feuilles nécessaires
df = pd.read_excel(fichier_data, sheet_name='SCHEMA_T')
df_comp = pd.read_excel(fichier_data, sheet_name='SCHEMA_COMPLEMENT')

if "N° SCENARIO" not in df.columns:
    print("Colonne 'N° SCENARIO' absente dans la feuille SCHEMA_T")
    sys.exit(1)

scenarios = list(df['N° SCENARIO'].dropna().unique())

for idx, num_scenario in enumerate(scenarios):
    df_schema = df[df['N° SCENARIO'] == num_scenario]
    try:
        ligne_Me = df_comp['Ligne ME'].dropna().iloc[0]
        Code_Regle = df_comp['Code Régle'].dropna().iloc[0]
        Libellé_schéma = df_comp["Libellé Schéma"].dropna().iloc[0]
        Type_du_schéma = df_comp["Type Schéma"].dropna().iloc[0]
        Num_Shema = df_comp["Num Schéma"].dropna().iloc[0]
        Condition_schéma = df_schema["CONDITION SCHEMA"].dropna().iloc[0] if "CONDITION SCHEMA" in df_schema else "N/A"
    except Exception as e:
        print(f"Erreur avec scenario {num_scenario} : {e}")
        continue

    # Ajouter paragraphes descriptifs
    for ligne in [
        f"Ligne de ME : {ligne_Me}",
        f"N° Schéma : {num_scenario}_{Code_Regle}-{Num_Shema}",
        f"Libellé du schéma : {Libellé_schéma}",
        f"Type du schéma : {Type_du_schéma}",
    ]:
        para = doc.add_paragraph()
        para.add_run(ligne)
    para_cond = doc.add_paragraph()
    para_cond.add_run(f"Condition du schema : {Condition_schéma}")
    colorer_texte(para_cond)

    # Génération des tableaux
    data = list(df_schema.iterrows())
    for i in range(0, len(data), 2):
        row_table = doc.add_table(rows=1, cols=2)
        row_table.autofit = True
        for j in range(2):
            if i + j >= len(data):
                continue
            _, row = data[i + j]
            try:
                condition_ME = row['CONDITION ME']
                compte_client = row['COMPTE']
                montant = row['MONTANT']
                sens = str(row['SENS']).replace('"', '').strip()
                cell = row_table.cell(0, j)

                p1 = cell.paragraphs[0]
                p1.add_run("")

                p2 = cell.add_paragraph()
                p2.alignment = 1
                run = p2.add_run(str(condition_ME))
                run.add_break()
                run = p2.add_run(str(compte_client))
                run.bold = True
                run.font.size = Pt(12)

                inner_table = cell.add_table(rows=1, cols=2)
                inner_table.autofit = True

                if sens == 'C':
                    inner_table.cell(0, 1).text = str(montant)
                elif sens == 'D':
                    inner_table.cell(0, 0).text = str(montant)

                set_cell_border(inner_table.cell(0, 0), top={"sz": 24, "val": "single", "color": "000000"})
                set_cell_border(inner_table.cell(0, 1), top={"sz": 24, "val": "single", "color": "000000"}, left={"sz": 24, "val": "single", "color": "000000"})

                traiter_cellules(inner_table)

            except KeyError as e:
                print(f"Colonne manquante : {e}")
                continue

    doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    # Saut de page toutes les 2 itérations (scénarios)
    if (idx + 1) % 2 == 0 and (idx + 1) < len(scenarios):
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# Appliquer la mise en forme globale Arial, taille 9
for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = 'Arial'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        run.font.size = Pt(9)

for table in doc.tables:
    traiter_cellules(table)

# Sauvegarder le document final
doc.save(fichier_output)
print("Document genere et colore avec succes :", fichier_output)
